import PySimpleGUI as sg
import sqlite3 as sql
from random import randint as r
from time import sleep
import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

global session_data
global wind_main
global wind_log

class Docs():
    @staticmethod
    def writecertificate(dest):
        doc = docx.Document()
        pars = []
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run('МИНОБРНАУКИ РОССИИ')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run('Федеральное государственное бюджетное образовательное учреждение высшего образования')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run('«Чувашский государственный университет имени И. Н. Ульянова»')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run('(ФГБОУ ВО «ЧГУ имени И. Н. Ульянова»)')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run('428015, г. Чебоксары, Московский проспект, 15 ')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        d = datetime.date.today()
        run = pars[len(pars)-1].add_run(f'{d.year}г.')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pars[len(pars)-1].add_run(f'СПРАВКА № {session_data[3]}')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.bold = True
        pars.append(doc.add_paragraph())
        pars[len(pars)-1].alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
        run = pars[len(pars)-1].add_run(f"Выдана в том, что {session_data[2]}, {session_data[4]} г.р. действительно является обучающимся {session_data[5]} формы обучения {session_data[6]} курса факультета {session_data[7]} направления подготовки (специальности) {session_data[8]} на {session_data[9]} основе. ")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pars.append(doc.add_paragraph())
        run = pars[len(pars)-1].add_run(f"Приказ о зачислении № {session_data[10]} ст от {session_data[11]} г.")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pars.append(doc.add_paragraph())
        run = pars[len(pars)-1].add_run("Выдана для предоставления по месту требования.")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        pars.append(doc.add_paragraph())
        pars.append(doc.add_paragraph())
        run = pars[len(pars)-1].add_run("Декан факультета___Щипцова")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.italic = True
        doc.save(dest +'Справка_'+ str(session_data[2]) +'.docx')

class DataBase():
    def __init__(self, conpath: str):
        self.db = sql.connect(conpath)
        self.curs = self.db.cursor()
        if self.db == None:
            print('Не получилось подклчиться!')
        else:
            print('Успешное подключение!')
            
    def loginCheck(self,login, password) -> bool:
        if login == '' or password == '':
            return False
        self.curs.execute(f""" SELECT login FROM students WHERE students.login = '{login}' AND students.pass = '{password}' """)
        logs = self.curs.fetchone()
        if logs == None:
            return False
        else:
            return True
    
    def regQuery(self, login, password, alias, datebirth, form, grade,faculty, spec, option,doe): # надо переделать
        decreeNo = r(1000,5000)
        self.curs.execute(f"INSERT INTO students(id, login,pass, initials, birthdate, formstud, grade, faculty, speciality, chargingOption, dayOfEnrollment, decreeNo, hascertificate) VALUES (NULL, ?,?,?,?,?,?,?,?,?,?,?,-1)",
        [login, password, alias, datebirth,form, grade,faculty, spec, option, doe, decreeNo])
        self.db.commit()
    
    def giveSessionInfo(self,login):
        self.curs.execute(f"SELECT id,login,initials,hascertificate, birthdate,formstud,grade,faculty,speciality,chargingOption, decreeNo, dayOfEnrollment FROM students WHERE students.login = ?", [login])
        return self.curs.fetchone()

    def getQueries(self): #История запросов
        self.curs.execute(f"SELECT * FROM queries WHERE who_id = ?",[session_data[0]])
        res = list(self.curs.fetchall())
        if res == []:
            return ['Пусто']
        else:
            return res
    
    def sendQuerry(self):
        self.curs.execute("SELECT id,sign FROM deanery WHERE initials = 'Щипцова'")
        dean = self.curs.fetchone()
        print(dean)
        self.curs.execute(f"INSERT INTO queries VALUES (NULL, -1,?,?,?)", [dean[1],session_data[0],dean[0]])
        self.db.commit()
    def approveQuerry(self):
        self.curs.execute(f"SELECT id FROM queries WHERE who_id = ? AND isApproved = -1",[session_data[0]])
        res = self.curs.fetchone()
        self.curs.execute(f"UPDATE students SET hascertificate = ? WHERE id = ?",[res[0],session_data[0]])
        self.db.commit()
        self.curs.execute(f"UPDATE queries SET isApproved = 1 WHERE id = ? AND who_id = ?",[res[0],session_data[0]])
        self.curs.execute(f"SELECT * FROM queries WHERE who_id = ? AND isApproved = -1",[session_data[0]])
        res = self.curs.fetchone()
        print(res)
        self.db.commit()
    def giveQueryInfo(self,id):
        self.curs.execute(f"SELECT * FROM queries WHERE id = ?", [id])    
        return self.curs.fetchone()
     
def window_reg():
    layout = [
        [sg.Text('Регистрация нового студента')],
        [sg.Text('Логин', size=(10,1)), sg.Input(key='-LOGREG-')],
        [sg.Text('Пароль',size=(10,1)), sg.Input(key='-PASSREG-')],
        [sg.Text('ФИО',size=(10,1)), sg.Input(key='-ALIAS-')],
        [sg.Text('Дата рождения'), sg.Input(key='-BIRTH-')],
        [sg.Text('Форма обучения'), sg.Combo(['Очная', 'Заочная', 'Очно-заочная'], key='-COMBFORM-')],
        [sg.Text('Курс',size=(10,1)), sg.Input(key='-GRADE-', size=(2,1))],
        [sg.Text('Факультет',size=(10,1)), sg.Input(key='-FACULTY-')],
        [sg.Text('Специальность'), sg.Input(key='-SPEC-')],
        [sg.Text('Основа',size=(10,1)), sg.Combo(['Бюджетная', 'Платная'], key='-COMBCHARGE-')],
        [sg.Text('Дата зачисления'), sg.Input(key='-DOE-')],
        [sg.Button('Принять',size=(10,1)), sg.Button('Выйти')]
    ]
    
    return sg.Window('Регистрация', layout, size=(300,350))

def window_login():
    layout = [
        [sg.Text('Введите логин и пароль')],
        [sg.Text('Логин',size=(6,1)), sg.Input(key='-LOGIN-')],
        [sg.Text('Пароль',size=(6,1)), sg.Input(key='-PASSWORD-', password_char='*')],
        [sg.Button('Войти'), sg.Button('Зарегистрироваться')]
    ]
    return sg.Window('Представьтесь', layout)
    

def window_student():
    layout = [
        [
                [sg.Text(f'Вы успешно вошли как {session_data[2]}!')],
                [sg.Button('Справка'), sg.Button('Выйти'),sg.Button('Отправить запрос')]
        ],
        [
                [sg.Text('История запросов'), sg.Button('Обновить')],
                [sg.Listbox(db.getQueries(),key='-LISTSTUD-', size=(30,10))],
                [[sg.Button('Подробнее')]]
                          
        ],
            
    ]

    return sg.Window('Личный кабинет', layout, size=(275,310))

wind_log = window_login()
db = DataBase('C:\\Users\\1\\OneDrive\\Учёба\\Учёба\\ИнфСисИТехи\\Справка\\certificate.db')
while True:
    event, values = wind_log.read()
    if event == sg.WIN_CLOSED:
        break
    if event == 'Войти':
        log_ch = db.loginCheck(values['-LOGIN-'], values['-PASSWORD-'])
        if log_ch:
            session_data = db.giveSessionInfo(values['-LOGIN-'])
        else:
            sg.popup_error('Неверный логин или пароль!', title= 'Ошибка!')
            continue
        break
    if event == 'Зарегистрироваться':
        wind_reg = window_reg()
        while True:
            event1, values1 = wind_reg.read()
            if event1 == sg.WIN_CLOSED or event1 == 'Выйти':
                break
            if event1 == 'Принять':
                log_ch = db.loginCheck(values1['-LOGREG-'], values1['-PASSREG-'])
                if log_ch:
                    sg.popup_error('Данные логин и пароль уже существуют!', title= 'Ошибка!')
                    continue
                else:
                   db.regQuery(values1['-LOGREG-'], values1['-PASSREG-'] ,values1['-ALIAS-'], values1['-BIRTH-'],
                   wind_reg['-COMBFORM-'].get(),values1['-GRADE-'],values1['-FACULTY-'],values1['-SPEC-'],
                   wind_reg['-COMBCHARGE-'].get(),values1['-DOE-'])
                   sg.popup('Вы прошли регистрацию!', title='Успешно!')
                break             
        wind_reg.close()
wind_log.close()

window_main = window_student()
while True:
    event, values = window_main.read()
    if event == 'Выйти' or event == sg.WIN_CLOSED:
         break
    if event == 'Справка':
        if (session_data[3] == -1):
            sg.popup('У вас нет справки!', title='Справки нет!')
        else:
            if (sg.popup_yes_no('У вас есть справка. Хотите ли вы сохранить её на компьютере?',title='Сохранить справку?')) == 'Yes':
                dest = sg.popup_get_folder('Укажите место сохранения','Место сохранения', 'D:\\')  
                Docs.writecertificate(dest)
                sg.popup('Справка успешно сохранена!', title='Сохранено!')
    if event == 'Обновить':
        upd = db.getQueries()
        session_data = db.giveSessionInfo(session_data[1])
        window_main['-LISTSTUD-'].update(upd)
    if event == 'Отправить запрос':
        if (session_data[3] != -1):
            if sg.popup_yes_no('У вас уже есть справка! Хотите ли создать новую?', title='Справка уже есть!') == 'No':
                continue
        db.sendQuerry()
        db.approveQuerry()
        sleep(1)
        sg.popup('Справка одобрена!')
    if event == 'Подробнее':
        try:
            id_q = window_main['-LISTSTUD-'].get()[0][0]
        except:
            sg.popup_error('Ничего не выбрано!', title='Ошибка!')
            continue
        query = list(db.giveQueryInfo(id_q))
        if (query[1] == 1):
            query[1] = 'Да'
        elif (query[1] == 0):
            query[1] == 'Нет'
        sg.popup(['ID запроса: ' + str(query[0]),'Одобрено: ' + str(query[1]), 'Подпись: ' + str(query[2]),
                  'ID студента: ' + str(query[3]), 'ID декана: ' + str(query[4])] , title='Подробнее по запросу ' + str(id_q))
        
    
    
window_main.close()
